import io
import re
from docx import Document

def parse_docx_to_html(file_buffer):
    """Convert DOCX buffer to HTML with placeholders"""
    try:
        doc = Document(io.BytesIO(file_buffer))
        html_parts = []
        all_fields = []
        seen_fields = set()
        
        html_parts.append('''
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            line-height: 1.6;
            padding: 40px;
            max-width: 900px;
            margin: 0 auto;
            background: white;
        }
        h1 { font-size: 18pt; text-align: center; color: #1e2d4a; margin: 20px 0; }
        h2 { font-size: 16pt; color: #1e2d4a; margin: 15px 0 10px; }
        h3 { font-size: 14pt; color: #1e2d4a; margin: 12px 0 8px; }
        p { margin-bottom: 10px; }
        ul, ol { margin: 8px 0 12px 24px; }
        li { margin-bottom: 4px; }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 15px 0;
        }
        th, td {
            border: 1px solid #cbd5e1;
            padding: 8px 12px;
            text-align: left;
            vertical-align: top;
        }
        th {
            background: #f8fafc;
            font-weight: bold;
        }
        .doc-field {
            display: inline-block;
            min-width: 150px;
            padding: 4px 8px;
            font-family: 'Times New Roman', Times, serif;
            font-size: 12pt;
            background: #fefce8;
            border: none;
            border-bottom: 2px solid #fbbf24;
            outline: none;
            transition: all 0.2s;
        }
        .doc-field:focus {
            background: #eff6ff;
            border-bottom-color: #2e6fea;
        }
        .signature-section {
            margin-top: 40px;
            padding: 20px;
            background: #f8fafc;
            border-left: 4px solid #2e6fea;
        }
    </style>
</head>
<body>
''')
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if heading
                style = para.style.name if para.style else ''
                
                if 'Heading 1' in style:
                    html_parts.append(f'<h1>{para.text}</h1>')
                elif 'Heading 2' in style:
                    html_parts.append(f'<h2>{para.text}</h2>')
                elif 'Heading 3' in style:
                    html_parts.append(f'<h3>{para.text}</h3>')
                else:
                    # Check for placeholders
                    text = para.text
                    fields = find_placeholders(text)
                    
                    for field in fields:
                        if field['name'] not in seen_fields:
                            seen_fields.add(field['name'])
                            all_fields.append(field)
                    
                    html_parts.append(f'<p>{text}</p>')
        
        # Process tables
        for table in doc.tables:
            html_parts.append('<table>')
            for i, row in enumerate(table.rows):
                tag = 'th' if i == 0 else 'td'
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<{tag}>{cell.text}</{tag}>')
                html_parts.append('</tr>')
            html_parts.append('</table>')
        
        html_parts.append('</body>\n</html>')
        
        html = ''.join(html_parts)
        
        # Replace placeholders with input fields
        for field in all_fields:
            placeholder = field['label']
            input_html = f'<input type="text" class="doc-field" data-field="{field["name"]}" placeholder="{placeholder}">'
            html = html.replace(f'[{placeholder}]', input_html)
            html = html.replace(f'{{{{{placeholder}}}}}', input_html)
        
        return {
            'success': True,
            'html': html,
            'fields': all_fields,
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables)
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def find_placeholders(text):
    """Find placeholders like [Field Name] or {{Field Name}}"""
    placeholders = []
    
    # Find [Field Name]
    bracket = re.findall(r'\[([^\]]+)\]', text)
    # Find {{Field Name}}
    curly = re.findall(r'\{\{([^}]+)\}\}', text)
    
    all_fields = bracket + curly
    
    for field in all_fields:
        field_clean = field.strip()
        field_name = field_clean.lower().replace(' ', '_').replace('-', '_')
        placeholders.append({
            'name': field_name,
            'label': field_clean,
            'type': 'date' if 'date' in field_name else 'text'
        })
    
    return placeholders